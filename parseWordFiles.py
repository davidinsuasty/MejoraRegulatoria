# -*- coding: utf-8 -*-
"""
Created on Tue May 19 21:04:11 2020

@author: hinsuasti

script donde se encuentran las funcionalidades para leer los archivos word y 
devolver los textos planos separadados para su posterior analisis
"""

from docx import Document
from itertools import repeat
from load_stopwords import load_stopwords
from limpieza_texto import limpieza_texto, limpieza_basica
import glob
import parmap
import pickle
import re
import time
import numpy as np


def splitWordFile2(path=None):
    """Función que separa el documento word en una lista de textos
    planos por cada tipo de ley, decreto, resolución, entre otros:

    :param str path: Ruta completa en donde está guardado el archivo *.dox

    :returns: Lista de textos planos separados por tipo de ley, decreto, entre otros.
    :rtype: list
    """
    name_file = path.split("\\")[-1][:-5]
    doc = Document(path)

    # Variables locales
    document_id, types, years, texts = [], [], [], []
    leyendo = False
    strings = ""

    # Lectura de las categorias que no tendrán en cuenta en los análisis
    # Para el análisis se tendrá en cuenta resoluciones, decretos y leyes
    ctgr = open(".\categorias_1.txt", encoding="utf-8", errors="ignore").read()
    ctgr = ctgr.lower().split("\n")
    condition = f"^({ctgr[0]}|{ctgr[1]}|{ctgr[2]}|{ctgr[3]}|{ctgr[4]}|{ctgr[5]}|{ctgr[6]}|{ctgr[7]}|{ctgr[8]}|{ctgr[9]}|{ctgr[10]}).*(20[012][0-9])$"
    # Este bucle recorre todo el texto línea a línea
    for p in doc.paragraphs:
        text = p.text.strip(" \t\n\r")
        if len(text) < 3:
            continue
        # remover espacios multiples
        text = re.sub(r" +", " ", text)
        if leyendo:
            if re.sub(r" +", "", text) == "***":
                texts.append(strings)
                strings = ""
                leyendo = False
                continue
            else:
                strings += " " + text
                continue
        if not leyendo:
            esNorma = re.findall(condition, text.lower())
            if esNorma:
                document_id.append(name_file)
                years.append(text[-4:])
                types.append(text.lower())
                leyendo = True
    if len(strings) > 0:
        texts.append(strings)

    if len(texts) != len(types):
        print("---->>>  Errores en documento -> {0}".format(path))

    return document_id, types, years, texts


def splitWordFile(path=None):
    """Función que separa el documento word en una lista de textos
    planos por cada tipo de ley, decreto, resolución, entre otros:

    :param str path: Ruta completa en donde está guardado el archivo *.dox

    :returns: Lista de textos planos separados por tipo de ley, decreto, entre otros.
    :rtype: list
    """
    name_file = path.split("\\")[-1][:-5]
    doc = Document(path)

    # Variables locales
    document_id, types, years, texts = [], [], [], []
    control_archivos, conteo = 0, 0
    strings = ""

    # Lectura de las categorias que no tendrán en cuenta en los análisis
    # Para el análisis se tendrá en cuenta resoluciones, decretos y leyes
    categorias = open(
        ".\categorias_1.txt", encoding="utf-8", errors="ignore"
    ).read()
    categorias = set(categorias.lower().split("\n"))

    # Este bucle recorre todo el texto línea a línea
    for p in doc.paragraphs:
        text = p.text
        if len(p.text.strip()) < 1:
            continue
        if p.style.name == "Heading 3":
            if p.text.strip() == "* * *" or p.text.strip() == "***":
                if conteo == 1:
                    texts.append(strings)
                    strings = ""
                control_archivos, conteo = 0, 0
                continue

            listacat = set(p.text.lower().split(" "))
            interseccion = listacat & categorias

            if len(interseccion) > 0:
                document_id.append(name_file)
                # match = re.match(r'.*([1-3][0-9]{3})', p.text)
                years.append(p.text.strip()[-4:])
                types.append(p.text.strip())
                control_archivos, conteo = 1, 1

        elif p.style.name != "Heading 3":
            if p.text.strip() == "* * *" or p.text.strip() == "***":
                if conteo == 1:
                    texts.append(strings)
                    strings = ""
                control_archivos, conteo = 0, 0
                continue
            if control_archivos == 1:
                strings = strings + " " + p.text
    if len(strings) > 0:
        texts.append(strings)

    if len(texts) != len(types):
        print("---->>>  Errores en documento -> {0}".format(path), end=" ")

    return document_id, types, years, texts


def load_docx_files(directorios):
    document_id_all = []
    text_all = []
    types_all = []
    years_all = []
    cont = 0
    n_docs = len(directorios)
    for doc_file in directorios:
        cont += 1
        print(
            "{0} of {1} -> file: {2}".format(cont, n_docs, doc_file), end=" "
        )
        document_id, types, years, texts = splitWordFile2(doc_file)
        document_id_all += document_id
        types_all += types
        years_all += years
        text_all += texts
        print(" -- {0} normas encontradas".format(len(types)))
    print("Total de normas para procesar: ", len(types_all))
    return document_id_all, types_all, years_all, text_all


def get_paths(path):
    folders = [f.split("\\")[4] for f in glob.glob(path + "/*/")]
    folders_with_enlace = [
        f.split("\\")[4] for f in glob.glob(path + "/*/[3-9]*enlace.docx")
    ]
    filepaths = glob.glob(path + "/*/[3-9]*enlace.docx")
    others = []
    folders_no_enlace = list(set(folders).difference(folders_with_enlace))
    for f in folders_no_enlace:
        full_path = path + f + "/"
        files = glob.glob(full_path + "[3-9]*.docx")
        max_version = 0
        for f_i in files:
            crr_version = re.findall(r"\d+", f_i)
            if len(crr_version) > 3:
                if max_version < int(crr_version[-1:][0]):
                    max_version = int(crr_version[-1:][0])
                    path_maxversion = f_i

        if max_version > 0:
            filepaths.append(path_maxversion)
        else:
            try:
                others.append(glob.glob(path + f + "/[3-9]*final.pdf")[0])
            except:
                try:
                    others.append(glob.glob(path + f + "/[3-9]*[1-9].pdf")[0])
                except:
                    continue

            # others.append(-1)

    return filepaths, others


# use example
if __name__ == "__main__":

    # path de los documentos word
    anios = list(range(2014, 2023))
    for anio in anios:
        print("+" + 50 * "-" + "+")
        print(
            "+" + 14 * " " + "procesando año -> " + str(anio) + 14 * " " + "+"
        )
        print("+" + 50 * "-" + "+")
        path = "..\..\Insumos\diario_oficial/" + str(anio) + "/"
        filepaths, pdffilepaths = get_paths(path)
        print("parsing word files...")
        doc_ids, types, year, texts = load_docx_files(filepaths)
        print("Done!")
        print("Limpieza basica...")
        texts2 = parmap.map(
            limpieza_basica,
            texts,
            pm_parallel=True,
            pm_pbar=True,
            pm_processes=12,
        )
        # terminos vinculantes
        # terms_vinc, n_palabras = get_terminos_vinculantes(texts)
        #%% limpieza de texto
        # cargar stopwords
        lp, le = load_stopwords(".\listas_stopwords")

        # remover stopwords
        print("Cleaning text...")
        cleanText = parmap.starmap(
            limpieza_texto,
            list(zip(texts, repeat(lp), repeat(le))),
            pm_parallel=True,
            pm_pbar=True,
            pm_processes=12,
        )

        data = {
            "data": cleanText,
            "texts": texts,
            "texts2": texts2,
            "ids": doc_ids,
            "type": types,
            "year": year,
        }

        ts = time.strftime("%m-%d-%Y_%H%M", time.localtime())
        pickle.dump(
            data,
            open(".\\data\\" + str(anio) + "-data_all.pkl", "wb"),
        )
